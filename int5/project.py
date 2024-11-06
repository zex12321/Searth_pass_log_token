import os
import json
import random
import string
import pandas as pd
from docx import Document 
from tqdm import tqdm  

# Убедимся, что файл создаётся в нужной директории
os.makedirs('leak_data', exist_ok=True)

# Функция для генерации случайной строки
def random_string(length):
    letters = string.ascii_letters + string.digits
    return ''.join(random.choice(letters) for _ in range(length))

# Генерация данных с индикатором прогресса
def generate_data(num_records=100000):
    domains = ['@bk.ru', '@mail.ru', '@yandex.ru', '@gmail.com']
    data_records = []
    
    # Используем tqdm для отображения прогресса
    for _ in tqdm(range(num_records), desc="Генерация данных"):
        data_records.append({
            'login': random_string(8) + random.choice(domains),
            'password': random_string(12),
            'token': random_string(16),
        })
    
    return data_records

# Генерация данных
data_records = generate_data()

# Оптимизированная запись в текстовый файл с использованием чанков
with open('leak_data/leak_text.txt', 'w') as f:
    for record in tqdm(data_records, desc="Запись в текстовый файл"):
        f.write(f"login: {record['login']}, password: {record['password']}, token: {record['token']}\n")

# Генерация CSV файла с помощью tqdm
with tqdm(total=len(data_records), desc="Запись в CSV файл") as pbar:
    df = pd.DataFrame(data_records)
    df.to_csv('leak_data/leak_data.csv', index=False)
    pbar.update(len(data_records))

# Генерация JSON файла с помощью tqdm
with open('leak_data/leak_data.json', 'w') as f:
    with tqdm(total=len(data_records), desc="Запись в JSON файл") as pbar:
        json.dump(data_records, f)
        pbar.update(len(data_records))

# Генерация Excel файла с помощью tqdm
with tqdm(total=len(data_records), desc="Запись в Excel файл") as pbar:
    df.to_excel('leak_data/leak_data.xlsx', index=False)
    pbar.update(len(data_records))

# Генерация DOCX файла (временно закомментировано)
# doc = Document()
# doc.add_heading('Случайные данные', level=1)
# for record in tqdm(data_records, desc="Запись в DOCX файл"):
#     doc.add_paragraph(f"login: {record['login']}, password: {record['password']}, token: {record['token']}")
# doc.save('leak_data/leak_data.docx')

# Генерация MD файла с индикатором прогресса
with open('leak_data/leak_data.md', 'w') as md_file:
    md_file.write('# Случайные данные\n\n')
    for record in tqdm(data_records, desc="Запись в MD файл"):
        md_file.write(f"login: {record['login']}, password: {record['password']}, token: {record['token']}\n")
